# -*- coding: utf-8 -*-
"""C3: convert references + in-text citations to APA 7th style, and fix Figure 2 RNH1 GTEx |Z|.
In-text citations are converted ONLY in the body (before References heading) to avoid
touching the reference-list author segments ("et al. YEAR.").
"""
import re
from docx import Document

SRC = "E:/workbuddy/2026-07-19-16-45-02/corrected/manuscript_GeneticEpidemiology_v9_final.docx"
OUT = "E:/workbuddy/2026-07-19-16-45-02/corrected/manuscript_GeneticEpidemiology_v9_final.docx"  # overwrite final

doc = Document(SRC)

# locate headings
ref_idx = None
fl_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "References":
        ref_idx = i
    if p.text.strip() == "Figure Legends":
        fl_idx = i
print("ref_idx=%s fl_idx=%s" % (ref_idx, fl_idx))

# ---------- In-text citation normalization (body only) ----------
def setp(p, new):
    for r in p.runs:
        r.text = ""
    if p.runs:
        p.runs[0].text = new
    else:
        p.add_run(new)

body_changed = 0
for i, p in enumerate(doc.paragraphs):
    if ref_idx is not None and i >= ref_idx:
        break
    t = p.text
    new = t
    # A: parenthetical citations "(X et al. YEAR)" -> "(X et al., YEAR)"
    new = re.sub(r'\(([A-Z][\w\-]+) et al\. (\d{4})\)', r'(\1 et al., \2)', new)
    # B: narrative "X et al. YEAR" (not inside parens) -> "X et al. (YEAR)"
    new = re.sub(r'(?<!\()([A-Z][\w\-]+) et al\. (\d{4})(?=[ ,.;])', r'\1 et al. (\2)', new)
    if new != t:
        setp(p, new)
        body_changed += 1
print("in-text citations normalized in %d body paragraphs" % body_changed)

# ---------- Reference list conversion (ref_idx+1 .. fl_idx-1) ----------
def conv_ref(s):
    s = re.sub(r'(\b\d{4})\. (?=")', r'(\1). ', s)          # year -> (Year) before straight quote
    s = re.sub(r'(\b\d{4})\. (?=\u201c)', r'(\1). ', s)     # year -> (Year) before curly quote
    s = re.sub(r', no\. (\d+):', r'(\1),', s)               # ", no. Issue:" -> "(Issue),"
    s = re.sub(r'([A-Z][\w\&\-\u2019 ]*?) (\d+)\(', r'\1, \2(', s)  # comma after journal name
    s = re.sub(r'(\d+): (?=\d)', r'\1, ', s)                # "Vol: Pages" -> "Vol, Pages"
    s = re.sub(r'(https://doi\.org/[^\s]+)\.', r'\1', s)    # strip trailing period after DOI
    s = re.sub(r'\.{2,}', '.', s)                           # collapse double periods (e.g. "Y..")
    if 'bioRxiv.' in s and 'doi.org' not in s:              # Gerlach 2025 preprint DOI
        s = s.replace('bioRxiv.', 'bioRxiv. https://doi.org/10.64898/2025.12.19.695550')
    return s

ref_changed = 0
for i, p in enumerate(doc.paragraphs):
    if ref_idx is None or fl_idx is None:
        continue
    if ref_idx < i < fl_idx:
        new = conv_ref(p.text)
        if new != p.text:
            setp(p, new)
            ref_changed += 1
print("reference entries converted: %d" % ref_changed)

# ---------- Figure 2 caption: RNH1 GTEx |Z|=8.5 -> 13.8 ----------
fig2 = 0
for p in doc.paragraphs:
    if "GTEx |Z|=8.5 to eQTLGen |Z|=11.6" in p.text:
        setp(p, p.text.replace("GTEx |Z|=8.5 to eQTLGen |Z|=11.6", "GTEx |Z|=13.8 to eQTLGen |Z|=11.6"))
        fig2 += 1
        print("Figure 2 caption RNH1 GTEx |Z| fixed")
print("fig2 fixed:", fig2)

doc.save(OUT)
print("SAVED ->", OUT)
