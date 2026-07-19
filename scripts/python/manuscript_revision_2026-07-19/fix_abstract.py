#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Restore the dropped abstract body in v9 and apply 40.7% -> 44.4%-51.9% correction.
Faithfully reconstructs v8's abstract structure: stray sentence -> Abstract heading -> body.
"""
import copy, shutil
from docx import Document
from docx.oxml.ns import qn

V9 = r"E:/workbuddy/2026-07-19-16-45-02/corrected/manuscript_GeneticEpidemiology_v9_final.docx"
V8 = r"E:/workbuddy/2026-07-19-16-45-02/corrected/manuscript_GeneticEpidemiology_v8_reconstructed.docx"
BAK = r"E:/workbuddy/2026-07-19-16-45-02/corrected/manuscript_GeneticEpidemiology_v9_final_BAK_before_abstract_restore.docx"

# backup
shutil.copy(V9, BAK)
print("backup ->", BAK)

# get exact abstract body from v8
v8 = Document(V8)
v8body = next(p.text for p in v8.paragraphs if "We present a dual-source eQTL sensitivity framework" in p.text)
assert "40.7% gene-level" in v8body, "expected 40.7% in v8 body"
v8body_fixed = v8body.replace("GTEx 40.7% gene-level", "GTEx 44.4%\u201351.9% per phenotype")
assert "40.7" not in v8body_fixed, "40.7 must be gone"
assert "44.4%" in v8body_fixed and "51.9%" in v8body_fixed
print("v8 body words:", len(v8body.split()), "| fixed words:", len(v8body_fixed.split()))

# open v9
doc = Document(V9)
paras = doc.paragraphs
# locate Abstract heading
heading_idx = next(i for i, p in enumerate(paras) if p.text.strip() == "Abstract")
print("Abstract heading at para", heading_idx)

# verify current state: para before heading should be stray sentence, para after should be empty
before = paras[heading_idx-1].text[:60]
after = paras[heading_idx+1].text
print("before heading:", repr(before))
print("after heading (should be empty):", repr(after[:40]))

# build the new body paragraph element by deep-copying v8's abstract paragraph (preserves style/runs)
v8doc = Document(V8)
v8_abs_para = next(p for p in v8doc.paragraphs if "We present a dual-source" in p.text)
new_el = copy.deepcopy(v8_abs_para._p)
# apply text correction inside the copied element's runs
for r in new_el.findall(qn('w:r')):
    t = r.find(qn('w:t'))
    if t is not None and t.text:
        if "GTEx 40.7% gene-level" in t.text:
            t.text = t.text.replace("GTEx 40.7% gene-level", "GTEx 44.4%\u201351.9% per phenotype")

# insert after heading
heading_el = doc.paragraphs[heading_idx]._p
parent = heading_el.getparent()
parent.insert(parent.index(heading_el) + 1, new_el)
print("inserted abstract body after heading")

# remove the now-redundant empty paragraph(s) right after the inserted body (keeps layout clean)
# find the newly inserted element's index
inserted_idx = parent.index(new_el)
# remove any immediately following empty <w:p> up to Keywords
to_remove = []
for sib in list(parent):
    # only remove empties between inserted body and the Keywords paragraph
    pass
# Simpler: remove a single following empty paragraph if present
nxt = new_el.getnext()
if nxt is not None:
    txt = "".join(nx.text or "" for nx in nxt.iter(qn('w:t')))
    if txt.strip() == "":
        parent.remove(nxt)
        print("removed trailing empty paragraph after abstract body")

doc.save(V9)
print("SAVED ->", V9)

# verify
doc2 = Document(V9)
p2 = [p.text for p in doc2.paragraphs]
hi = next(i for i, p in enumerate(p2) if p.strip() == "Abstract")
print("\nVERIFY: Abstract heading at", hi)
print("  para before:", repr(p2[hi-1][:50]))
print("  para after (body):", repr(p2[hi+1][:90]))
print("  body has 44.4%:", "44.4%" in p2[hi+1], "| has 51.9%:", "51.9%" in p2[hi+1], "| no 40.7%:", "40.7" not in p2[hi+1])
print("  body words:", len(p2[hi+1].split()))
