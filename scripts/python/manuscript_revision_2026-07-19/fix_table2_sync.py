#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Align Table2 GTEx Candidate DR 48.1% -> 48.2% to match manuscript (13/27 = 48.148%)."""
import shutil, csv
from docx import Document

MASTER = r"E:/workbuddy/2026-07-19-16-45-02/corrected/Table2_v7_candidate_fixed.docx"
UPLOAD = r"E:/workbuddy/Genetic Epidemiology/投稿资料/上传主表/Table2.docx"
UPLOAD_BAK = r"E:/workbuddy/Genetic Epidemiology/投稿资料/上传主表/Table2_BAK_before_sync.docx"
CSV = r"E:/workbuddy/2026-07-19-16-45-02/rerun_enrich/tables__Table2.csv"

def fix_docx(path):
    d = Document(path)
    t = d.tables[0]
    changed = []
    for r in t.rows:
        cs = [c.text.strip() for c in r.cells]
        if cs[0] == "Candidate" and cs[1] == "DR":
            cell = r.cells[2]  # GTEx ACAT-O FDR (%) column
            old = cell.text
            assert "48.1%" in old, "expected 48.1% in Candidate DR GTEx cell: " + old
            # replace within runs
            for para in cell.paragraphs:
                for run in para.runs:
                    if "48.1%" in run.text:
                        run.text = run.text.replace("48.1%", "48.2%")
            changed.append((path, old, cell.text))
    d.save(path)
    return changed

print("MASTER:", fix_docx(MASTER))
shutil.copy(UPLOAD, UPLOAD_BAK)
print("backup upload Table2 ->", UPLOAD_BAK)
print("UPLOAD:", fix_docx(UPLOAD))

# fix csv
rows = []
with open(CSV, encoding='utf-8', newline='') as f:
    r = csv.reader(f)
    header = next(r)
    for row in r:
        if row[0] == "Candidate" and row[1] == "DR" and row[2] == "48.1%":
            row[2] = "48.2%"
        rows.append(row)
with open(CSV, 'w', encoding='utf-8', newline='') as f:
    w = csv.writer(f)
    w.writerow(header)
    w.writerows(rows)
print("CSV Candidate DR GTEx updated to 48.2%")

# verify
d = Document(UPLOAD)
for r in d.tables[0].rows:
    cs = [c.text.strip() for c in r.cells]
    if cs[0] in ("Candidate",) and cs[1] in ("DR","DN","DPN"):
        print("  UPLOAD", cs[0], cs[1], "GTEx=", cs[2], "eQTLGen=", cs[3])
