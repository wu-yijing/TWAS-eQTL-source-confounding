#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Diff + fix Table 3 (col1 stale 40.7%) and propagate to uploaded submission file.
Also dump Table1/Table2 for the sync report.
"""
import copy, shutil, difflib
from docx import Document

MASTER = r"E:/workbuddy/2026-07-19-16-45-02/corrected/Table3_v7_candidate_fixed.docx"
UPLOAD = r"E:/workbuddy/Genetic Epidemiology/投稿资料/上传主表/Table3.docx"
UPLOAD_BAK = r"E:/workbuddy/Genetic Epidemiology/投稿资料/上传主表/Table3_BAK_before_sync.docx"

def cells_of(path):
    d = Document(path)
    t = d.tables[0]
    return [[c.text.strip() for c in r.cells] for r in t.rows]

master_cells = cells_of(MASTER)
upload_cells = cells_of(UPLOAD)

print("=== FULL DIFF: MASTER vs UPLOAD Table3 ===")
diff_found = False
for ri in range(max(len(master_cells), len(upload_cells))):
    m = master_cells[ri] if ri < len(master_cells) else []
    u = upload_cells[ri] if ri < len(upload_cells) else []
    for ci in range(max(len(m), len(u))):
        mv = m[ci] if ci < len(m) else ""
        uv = u[ci] if ci < len(u) else ""
        if mv != uv:
            diff_found = True
            print(f"  ROW{ri} COL{ci}:")
            print(f"    MASTER : {mv}")
            print(f"    UPLOAD : {uv}")
print("  (no other diffs)" if not diff_found else "")

# ---- FIX MASTER col1 of 'Candidate gene FDR enrichment' row ----
NEW_COL1 = "48.2% (DR, 13/27), 51.9% (DN, 14/27), 44.4% (DPN, 12/27) FDR enrichment; comparable to non-candidate and T2DM-control groups"
OLD_COL1_FRAG = "40.7% cross-phenotype FDR enrichment (11/27 genes); highest among three groups"

d = Document(MASTER)
t = d.tables[0]
fixed = False
for r in t.rows:
    cs = [c.text.strip() for c in r.cells]
    if cs[0] == "Candidate gene FDR enrichment (GTEx vs eQTLGen)":
        cell = r.cells[1]
        full = cell.text
        assert OLD_COL1_FRAG in full, "frag not found: " + full
        new_full = full.replace(OLD_COL1_FRAG, NEW_COL1)
        # replace text in the cell's first paragraph (preserve formatting)
        para = cell.paragraphs[0]
        # clear runs, set combined text (cell has single para)
        # gather all run texts
        runs = para.runs
        if runs:
            # set first run to new_full, clear the rest
            runs[0].text = new_full
            for rr in runs[1:]:
                rr.text = ""
        else:
            para.add_run(new_full)
        # also fix any nested (merged) — ensure col1 updated
        fixed = True
        print("\nFIXED MASTER col1 ->", new_full[:90])
        break
assert fixed, "Candidate row not found in master"

d.save(MASTER)
print("SAVED master ->", MASTER)

# verify
d2 = Document(MASTER)
for r in d2.tables[0].rows:
    cs = [c.text.strip() for c in r.cells]
    if cs[0] == "Candidate gene FDR enrichment (GTEx vs eQTLGen)":
        print("  VERIFY master col1:", cs[1][:90])
        print("  VERIFY master col4 (post-match):", cs[4][:70])

# ---- PROPAGATE to upload (master is the corrected source of truth) ----
shutil.copy(UPLOAD, UPLOAD_BAK)
print("\nbackup upload ->", UPLOAD_BAK)
shutil.copy(MASTER, UPLOAD)
print("copied master -> upload:", UPLOAD)

# verify upload now matches master
uc = cells_of(UPLOAD)
mc = cells_of(MASTER)
match = all(uc[i][j] == mc[i][j] for i in range(len(uc)) for j in range(len(uc[i])))
print("UPLOAD now matches MASTER:", match)
