from docx import Document
import re, copy

OLD = "21238203"
NEW = "21428347"

def fix_docx(path, label):
    doc = Document(path)
    n = 0
    for p in doc.paragraphs:
        if OLD in p.text:
            # rebuild runs preserving formatting of first run
            full = p.text.replace(OLD, NEW)
            # collect run formatting from first run that had text
            ref_run = next((r for r in p.runs if r.text.strip()), None)
            for r in p.runs:
                r.text = ""
            if p.runs:
                p.runs[0].text = full
            else:
                # no runs: add one
                new_run = p.add_run(full)
            n += p.text.count(NEW)
    # also fix table cells
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                if OLD in c.text:
                    c.text = c.text.replace(OLD, NEW)
                    n += c.text.count(NEW)
    doc.save(path)
    print(f"{label}: {path}")
    print(f"  -> occurrences of {NEW} after fix: {n}")

fix_docx(r"E:/workbuddy/2026-07-19-16-45-02/corrected/manuscript_GeneticEpidemiology_v9_final.docx", "MANUSCRIPT v9")
fix_docx(r"E:/workbuddy/Genetic Epidemiology/投稿资料/cover_letter_GeneticEpidemiology.docx", "COVER LETTER")
