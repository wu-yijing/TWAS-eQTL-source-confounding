from docx import Document

SRC = r"E:/workbuddy/2026-07-19-16-45-02/corrected/manuscript_GeneticEpidemiology_v9_final.docx"
OUT = SRC

doc = Document(SRC)
body = doc.element.body
children = list(body)  # body-level elements (w:p, w:tbl, ...)

def elem_for_heading(text):
    for c in children:
        tag = c.tag.split('}')[-1]
        if tag == 'p':
            # get paragraph text
            texts = [t.text or '' for t in c.iter() if t.tag.split('}')[-1] == 't']
            if ''.join(texts).strip() == text:
                return c
    raise ValueError(f"heading not found: {text!r}")

methods_elem = elem_for_heading('Methods')
ac_elem = elem_for_heading('Author Contributions')
results_elem = elem_for_heading('Results')

# collect body children between methods_elem (inclusive) and ac_elem (exclusive)
block = []
capture = False
for c in children:
    if c is methods_elem:
        capture = True
    if c is ac_elem:
        break
    if capture:
        block.append(c)

print(f"Methods block: {len(block)} body elements")

# remove block
for el in block:
    body.remove(el)

# insert block before results_elem, preserving order
for el in block:
    results_elem.addprevious(el)

doc.save(OUT)
print("saved:", OUT)

# verify new order
doc2 = Document(OUT)
order = []
for p in doc2.paragraphs:
    t = p.text.strip()
    if t in ('Abstract','Introduction','Methods','Results','Discussion','Conclusions',
             'Author Contributions','Acknowledgements','Data availability','References','Figure Legends'):
        order.append(t)
print("NEW ORDER:", order)
