#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Final verification of manuscript_GeneticEpidemiology_v9_final.docx.
Checks residual false claims, correct values, and APA 7th formatting.
"""
import re
from docx import Document

SRC = r"E:/workbuddy/2026-07-19-16-45-02/corrected/manuscript_GeneticEpidemiology_v9_final.docx"

doc = Document(SRC)

# Collect full text by paragraph + tables
paras = [p.text for p in doc.paragraphs]
full = "\n".join(paras)

# Also gather table cell text
table_texts = []
for ti, t in enumerate(doc.tables):
    for r in t.rows:
        for c in r.cells:
            table_texts.append(c.text)
full_with_tables = full + "\n" + "\n".join(table_texts)

def count(pat, text):
    return len(re.findall(pat, text, flags=re.IGNORECASE))

print("=" * 70)
print("RESIDUAL FALSE-CLAIM SCAN (should all be 0)")
print("=" * 70)
checks = {
    "residual '40.7%' (gene-level false)": r"40\.7\s*%",
    "false '0%' eQTLGen enrichment claim": r"eqtlgen[^.\n]{0,30}0\s*%",
    "'Figure 6b' wrong ref": r"figure\s*6b",
    "Zenodo 21428347 (old ms DOI)": r"21428347",
    "Zenodo 21416787 (old cover DOI)": r"21416787",
    "Chicago ', no. ' journal pattern": r",\s*no\.\s*\d+:",
    "SMD<0.1 false caption": r"\|smd\|\s*<\s*0\.1",
    "RNH1 GTEx |Z| 8.5 (old)": r"\|z\|\s*=\s*8\.5",
}
for label, pat in checks.items():
    n = count(pat, full_with_tables)
    flag = "OK" if n == 0 else "!!! CHECK"
    print(f"  [{flag}] {label}: {n}")

print()
print("=" * 70)
print("REQUIRED CORRECT-VALUE PRESENCE (should all be >0)")
print("=" * 70)
present = {
    "Zenodo 21238203 (unified)": r"21238203",
    "GTEx DN 51.85% / 51.9%": r"51\.8?5\s*%",
    "GTEx DR 48.15% / 48.1%": r"48\.1?5?\s*%",
    "GTEx DPN 44.44% / 44.4%": r"44\.4?4?\s*%",
    "RNH1 GTEx Z DR 13.82/13.8": r"13\.8\s*2?",
    "RNH1 GTEx Z DN 7.00/7.0": r"7\.0\s*0?",
    "RNH1 GTEx Z DPN 8.57/8.6": r"8\.5\s*7?",
    "RNH1 eQTLGen Z +11.62": r"11\.62",
    "SMD -0.153 (body)": r"0\.153",
    "Figure 5b correct ref": r"figure\s*5b",
    "SMD<0.25 corrected caption": r"\|smd\|\s*<\s*0\.25",
}
for label, pat in present.items():
    n = count(pat, full_with_tables)
    flag = "OK" if n > 0 else "!!! MISSING"
    print(f"  [{flag}] {label}: {n}")

print()
print("=" * 70)
print("APA 7TH IN-TEXT CITATION FORMAT CHECK")
print("=" * 70)
# Narrative 'X et al. (YEAR)' and parenthetical '(X et al., YEAR)'
narr = count(r"[A-Z][a-z]+ et al\. \(\d{4}\)", full)
paren = count(r"\([A-Z][a-z]+ et al\., \d{4}\)", full)
paren_missing_comma = count(r"\([A-Z][a-z]+ et al\. \d{4}\)", full)
print(f"  Narrative 'X et al. (YEAR)': {narr}")
print(f"  Parenthetical '(X et al., YEAR)' with comma: {paren}")
print(f"  Parenthetical '(X et al. YEAR)' MISSING comma: {paren_missing_comma}")
flag = "OK" if paren_missing_comma == 0 else "!!! FIX NEEDED"
print(f"  [{flag}] parenthetical comma rule")

print()
print("=" * 70)
print("REFERENCES SECTION APA 7th SHAPE CHECK")
print("=" * 70)
# Find References heading
ref_start = None
for i, p in enumerate(paras):
    if re.match(r"^references$", p.strip(), flags=re.IGNORECASE):
        ref_start = i
        break
if ref_start is not None:
    refs = paras[ref_start+1:]
    refs = [r for r in refs if r.strip()]
    print(f"  References heading at para {ref_start}; {len(refs)} ref entries")
    # Check each ref has (Year) pattern and DOI
    no_year = [r for r in refs if not re.search(r"\(\d{4}\)", r)]
    no_doi = [r for r in refs if "doi.org" not in r and "doi:" not in r.lower() and "PMID" not in r]
    print(f"  Entries missing (Year): {len(no_year)}")
    for r in no_year[:10]:
        print(f"    - {r[:80]}")
    print(f"  Entries missing DOI/PMID: {len(no_doi)}")
    for r in no_doi[:10]:
        print(f"    - {r[:80]}")
else:
    print("  !!! References heading NOT found")

print()
print("=" * 70)
print("WORD COUNT of abstract (target <=200)")
print("=" * 70)
abs_start = None
for i, p in enumerate(paras):
    if re.match(r"^abstract$", p.strip(), flags=re.IGNORECASE):
        abs_start = i
        break
if abs_start is not None:
    # abstract may span until next heading
    words = []
    for p in paras[abs_start+1:abs_start+12]:
        if re.match(r"^(keywords|introduction|methods|results)", p.strip(), flags=re.IGNORECASE):
            break
        words.append(p)
    wc = len(" ".join(words).split())
    print(f"  Approx abstract word count: {wc}")
else:
    print("  !!! Abstract heading NOT found")

print()
print("DONE.")
